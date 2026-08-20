import docx
import io

class DOCXParser:
    @staticmethod
    def parse_docx(file_bytes: bytes) -> dict:
        """
        Parses DOCX file bytes to extract text, tables, headers/footers, and structural metrics.
        """
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            has_tables = len(doc.tables) > 0
            has_header_footer_text = False
            
            # Paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text.strip())
            
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            # Check headers / footers
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    for h_p in section.header.paragraphs:
                        if h_p.text.strip():
                            has_header_footer_text = True
                if section.footer and section.footer.paragraphs:
                    for f_p in section.footer.paragraphs:
                        if f_p.text.strip():
                            has_header_footer_text = True
                            
            raw_text = "\n".join(full_text)
            word_count = len(raw_text.split())
            
            # DOCX doesn't have explicit pages natively until rendered, estimate page count by ~450 words/page
            estimated_page_count = max(1, round(word_count / 450))
            
            return {
                "raw_text": raw_text,
                "page_count": estimated_page_count,
                "word_count": word_count,
                "has_images": False, # DOCX inline shape check optional
                "has_tables": has_tables,
                "has_multi_column": False,
                "has_header_footer_text": has_header_footer_text,
                "parser_confidence": 1.0 if word_count > 50 else 0.4,
                "success": True,
                "error": None
            }
        except Exception as e:
            return {
                "raw_text": "",
                "page_count": 0,
                "word_count": 0,
                "has_images": False,
                "has_tables": False,
                "has_multi_column": False,
                "has_header_footer_text": False,
                "parser_confidence": 0.0,
                "success": False,
                "error": str(e)
            }
